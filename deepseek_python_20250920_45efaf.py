import asyncio
import os
import httpx
import sqlite3
from datetime import datetime
from fastapi import APIRouter, WebSocket, WebSocketDisconnect
from bs4 import BeautifulSoup
import json
from pathlib import Path
import tempfile
import zipfile
from typing import List, Optional
from docx import Document
import aiofiles
from pydantic import BaseModel
from dotenv import load_dotenv

# Import MCP client
from fastmcp import Client
from fastmcp.client.transports import SSETransport

# MCP Server URL
MCP_SERVER_URL = "http://127.0.0.1:8002/sse"

# Global MCP client instance
mcp_client = None

async def get_mcp_client():
    """Get or create MCP client"""
    global mcp_client
    if mcp_client is None:
        transport = SSETransport(MCP_SERVER_URL)
        mcp_client = Client(transport)
        await mcp_client.__aenter__()
    return mcp_client

# ---------------------------
# Websocket manager
# ---------------------------
class ConnectionManager:
    def __init__(self):
        self.active_connections: set[WebSocket] = set()

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.add(websocket)

    def disconnect(self, websocket: WebSocket):
        self.active_connections.discard(websocket)

    async def broadcast(self, message: dict):
        for connection in self.active_connections.copy():
            try:
                await connection.send_json(message)
            except Exception:
                self.disconnect(connection)

manager = ConnectionManager()
router = APIRouter()

# ---------------------------
# Configuration
# ---------------------------
DB_PATH = r"C:\Users\342534\Desktop\Telecom Standards Management\backend\telecom_ai.db"
CONFIG_PATH = Path(r"C:\Users\342534\Desktop\Telecom Standards Management\backend\agents\config\crawler_config.json")

RECIPIENT_EMAILS = [
    "342534@nttdata.com",
    "harshitanagaraj.guled@nttdata.com",
    "neelambuz.singh@nttdata.com"
]

LATEST_STATUS = ""

def update_latest_status(msg: str):
    global LATEST_STATUS
    LATEST_STATUS = msg

def print_and_store(msg: str):
    update_latest_status(msg)
    try:
        asyncio.create_task(manager.broadcast({"type": "log", "data": msg}))
    except Exception:
        pass

async def broadcast_status():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT filename, url, status, last_checked FROM files ORDER BY last_checked DESC LIMIT 1")
    row = c.fetchone()
    conn.close()
    cfg = read_crawler_config()
    data = {
        "filename": row[0] if row else None,
        "url": row[1] if row else None,
        "status": row[2] if row else None,
        "last_checked": row[3] if row else None,
        "frequency": cfg["crawler_frequency"]
    }
    await manager.broadcast({"type": "status", "data": data})

# ---------------------------
# Config helpers
# ---------------------------
def read_crawler_config():
    if CONFIG_PATH.exists():
        with open(CONFIG_PATH, "r") as f:
            data = json.load(f)
            data.setdefault("crawler_active", False)
            data.setdefault("crawler_frequency", 10)
            data.setdefault("crawler_url", "https://www.3gpp.org/ftp/specs/archive/23_series/23.002")
            return data
    else:
        return {
            "crawler_active": False,
            "crawler_frequency": 10,
            "crawler_url": "https://www.3gpp.org/ftp/specs/archive/23_series/23.002"
        }

crawler_wakeup_event = asyncio.Event()

def write_crawler_config(active: bool = None, frequency: int = None, url: str = None):
    cfg = read_crawler_config()
    orig_active = cfg.get("crawler_active", False)
    if active is not None:
        cfg["crawler_active"] = active
    if frequency is not None:
        cfg["crawler_frequency"] = frequency
    if url is not None:
        cfg["crawler_url"] = url
    CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(CONFIG_PATH, "w") as f:
        json.dump(cfg, f)
    if active is not None and active != orig_active:
        try:
            crawler_wakeup_event.set()
        except Exception:
            pass

# ---------------------------
# Word document processing helpers
# ---------------------------
def extract_all_word_from_zip(zip_path: str, extract_dir: str) -> list[str]:
    word_paths = []
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        for file in zip_ref.namelist():
            if file.lower().endswith('.docx') or file.lower().endswith('.doc'):
                extracted_path = zip_ref.extract(file, extract_dir)
                word_paths.append(os.path.abspath(extracted_path))
    return word_paths

import docx2txt
import doc2txt

def extract_text_from_word(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        try:
            return docx2txt.process(file_path) or ""
        except Exception as e:
            print(f"Error extracting .docx: {e}")
            return ""
    elif ext == ".doc":
        try:
            return doc2txt.extract_text(file_path) or ""
        except Exception as e:
            print(f"Error extracting .doc: {e}")
            return ""
    else:
        return ""

def select_main_word(word_paths: list[str]) -> str | None:
    max_lines = 0
    main_word = None
    for path in word_paths:
        text = extract_text_from_word(path)
        lines = [line for line in text.splitlines() if line.strip()]
        if len(lines) > max_lines:
            max_lines = len(lines)
            main_word = path
    return main_word

import re

def clean_summary_text(text: str) -> str:
    lines = text.splitlines()
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append("")
            continue
        cleaned = re.sub(r"^(\*|\-|\+|\#)+\s*", "", stripped)
        cleaned_lines.append(cleaned)
    result = "\n".join(cleaned_lines)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()

def save_summary_to_docx(summary: str, comparison: str, filepath: str, doc_name: str) -> None:
    doc = Document()
    doc.add_heading(f"Summary for {doc_name}", 0)
    doc.add_heading("Overview (1-2 pages):", level=1)
    doc.add_paragraph(summary)
    doc.save(filepath)

# ---------------------------
# Database Helpers
# ---------------------------
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute(
        """CREATE TABLE IF NOT EXISTS files (
               id INTEGER PRIMARY KEY,
               filename TEXT UNIQUE,
               url TEXT,
               status TEXT,
               last_checked TEXT
           )"""
    )
    conn.commit()
    conn.close()

def add_file(filename: str, url: str, status: str):
    current_time = datetime.now().isoformat() 
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute(
        """INSERT INTO files (filename, url, status, last_checked)
           VALUES (?, ?, ?, ?)
           ON CONFLICT(filename) DO UPDATE SET
               status=excluded.status,
               last_checked=excluded.last_checked""",
        (filename, url, status, current_time)
    )
    conn.commit()
    conn.close()

def get_latest_file():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT filename, url, last_checked FROM files ORDER BY last_checked DESC LIMIT 1")
    row = c.fetchone()
    conn.close()
    return {"filename": row[0], "url": row[1], "last_checked": row[2]} if row else None

# ---------------------------
# MCP Tool Wrapper Functions
# ---------------------------
async def fetch_url_via_mcp(url: str) -> str:
    """Call MCP server to fetch URL"""
    client = await get_mcp_client()
    result = await client.call_tool("fetch_url", {"url": url})
    return result.data

async def parse_version_via_mcp(html: str) -> str:
    """Call MCP server to parse version"""
    client = await get_mcp_client()
    result = await client.call_tool("parse_version", {"html": html})
    return result.data

async def compare_versions_via_mcp(old: str, new: str) -> str:
    """Call MCP server to compare versions"""
    client = await get_mcp_client()
    result = await client.call_tool("compare_versions", {"old": old, "new": new})
    return result.data

async def should_crawl_reasoning_via_mcp(last_checked: str, last_file: str, frequency: int, current_time: str) -> bool:
    """Call MCP server for crawl decision"""
    client = await get_mcp_client()
    result = await client.call_tool("should_crawl_reasoning_llm", {
        "last_checked": last_checked,
        "last_file": last_file,
        "frequency": frequency,
        "current_time": current_time
    })
    return result.data

async def send_notification_via_mcp(to_email: str, subject: str, content: str, attachment_path: str = None) -> int:
    """Call MCP server to send notification"""
    client = await get_mcp_client()
    
    if attachment_path:
        result = await client.call_tool("send_notification", {
            "to_email": to_email,
            "subject": subject,
            "content": content,
            "attachment_path": attachment_path
        })
    else:
        result = await client.call_tool("send_notification", {
            "to_email": to_email,
            "subject": subject,
            "content": content
        })
    
    return result.data

# ---------------------------
# LLM Helper (keep this local)
# ---------------------------
from llm.llm_endpoints import chat_completion

# ---------------------------
# Agent Logic
# ---------------------------
async def monitor_site():
    url = read_crawler_config()["crawler_url"]
    
    # Use MCP tool instead of local function
    html = await fetch_url_via_mcp(url)
    await asyncio.sleep(2)
    
    # Use MCP tool for parsing
    latest_file = os.path.basename(await parse_version_via_mcp(html))
    
    if not latest_file:
        print(f"❌ No .zip files found at {url}")
        print_and_store(f"❌ No .zip files found at {url}")
        return

    last_seen = get_latest_file()
    last_filename = os.path.basename(last_seen["filename"]) if last_seen else None
    
    # Use MCP tool for comparison
    decision = await compare_versions_via_mcp(last_filename or "", latest_file)

    if not url.endswith('/'):
        url += '/'
    file_url = url + latest_file
    add_file(latest_file, url, decision)
    await broadcast_status()

    if decision == "new version":
        print(f"🚀 New file detected: {latest_file}")
        print_and_store(f"🚀 New file detected: {latest_file}")
        await asyncio.sleep(15)

        # Download and process the file
        with tempfile.TemporaryDirectory() as temp_dir:
            zip_path = os.path.join(temp_dir, latest_file)
            async with httpx.AsyncClient() as client:
                resp = await client.get(file_url)
                resp.raise_for_status()
                async with aiofiles.open(zip_path, 'wb') as f:
                    await f.write(resp.content)
            print_and_store(f"Downloaded ZIP to {zip_path}")

            word_paths = extract_all_word_from_zip(zip_path, temp_dir)
            if not word_paths:
                print_and_store("No .docx or .doc files found in ZIP.")
                return

            main_word_path = select_main_word(word_paths)
            if not main_word_path:
                print_and_store("Could not determine main Word file.")
                return

            print_and_store(f"Main Word file selected: {main_word_path}")

            main_text = extract_text_from_word(main_word_path)
            if not main_text.strip():
                print_and_store("No extractable text found in Word file.")
                return

            print_and_store("Generating summary...")
            await asyncio.sleep(5)

            chunk_size = 3000
            chunks = [main_text[i:i+chunk_size] for i in range(0, len(main_text), chunk_size)]
            summaries = []
            for idx, chunk in enumerate(chunks):
                summary = chat_completion(
                    user_prompt=f"Summarize the following telecom standard document content (chunk {idx+1}):\n\n{chunk}",
                    system_instruction="Summarize this technical content for a standards update digest. Focus on key changes, new features, and important scope."
                )
                summaries.append(summary)
            final_summary = "\n\n".join(summaries)
            print_and_store("Summary generated.")
            await asyncio.sleep(1.5)

            summary_docx_path = os.path.join(temp_dir, f"summary_{os.path.splitext(latest_file)[0]}.docx")
            cleaned_summary = clean_summary_text(final_summary)
            save_summary_to_docx(cleaned_summary, "", summary_docx_path, os.path.basename(main_word_path))
            print_and_store(f"Summary DOCX saved: {summary_docx_path}")

            subject = f"New 3GPP File Available: {latest_file}"
            content = (
                f"<p>Dear Team,</p>"
                f"<p>We are pleased to inform you that a new version of the 3GPP standard <b>{latest_file}</b> has been uploaded to the official 3GPP website.</p>"
                f"<p>Please find a brief summary of the document attached in this email, to help you quickly identify what is new.</p>"
                f"<p>Best regards,<br>Standards Monitoring Automation</p>"
            )

            print_and_store("Sending Email Notification with summary attached...")
            for email in RECIPIENT_EMAILS:
                status = await send_notification_via_mcp(email, subject, content, attachment_path=summary_docx_path)
                if status == 200:
                    print_and_store(f"Email sent successfully to {email}")
                else:
                    print_and_store(f"Failed to send email to {email}, status: {status}")
            print_and_store("Email Notification sent with summary attached!")

    else:
        print(f"No new file. Current latest: {latest_file}")
        print_and_store(f"No new file. Current latest: {latest_file}")

# ---------------------------
# Background Task
# ---------------------------
last_active = None

async def background_monitor():
    global last_active
    
    while True:
        cfg = read_crawler_config()
        freq = cfg["crawler_frequency"]
        if cfg["crawler_active"]:
            last = get_latest_file()
            last_checked = last['last_checked'] if last and last['last_checked'] else "never"
            last_file = last['filename'] if last else "none"
            
            # Use MCP tool for decision making
            now = datetime.now().isoformat()
            should_crawl = await should_crawl_reasoning_via_mcp(
                last_checked=last_checked,
                last_file=last_file,
                frequency=freq,
                current_time=now
            )
            
            if should_crawl:
                try:
                    await monitor_site()
                except Exception as e:
                    print(f"❌ Too many requests to website: {e}")
                    print_and_store(f"Website is loading....")
            else:
                print("Agentic Reasoning: Decided not to crawl this cycle.")
                print_and_store("Agentic Reasoning: Decided not to crawl this cycle.")
        else:
            if last_active != False:
                print("🔌 Crawler is inactive.")
                print_and_store("🔌 Crawler is inactive.")
        last_active = cfg["crawler_active"]

        try:
            crawler_wakeup_event.clear()
            await asyncio.wait_for(crawler_wakeup_event.wait(), timeout=freq)
            continue
        except asyncio.TimeoutError:
            pass

# ---------------------------
# FastAPI Endpoints (UNCHANGED)
# ---------------------------
@router.get("/monitor/agent")
async def get_agent_status():
    cfg = read_crawler_config()
    return {
        "active": cfg["crawler_active"],
        "frequency": cfg["crawler_frequency"],
    }

class AgentToggleRequest(BaseModel):
    active: bool

@router.post("/monitor/agent/toggle")
async def toggle_agent(body: AgentToggleRequest):
    write_crawler_config(active=body.active)
    msg = "🔌 Crawler activated." if body.active else "🔌 Crawler deactivated."
    print_and_store(msg)
    await broadcast_status()
    return {"success": True, "active": body.active}

class AgentUrlRequest(BaseModel):
    url: str

@router.post("/agent/{id}/url")
async def set_url(id: int, payload: dict):
    new_url = payload.get("url")
    write_crawler_config(url=new_url)
    await broadcast_status()
    return {"success": True}

class AgentFrequencyRequest(BaseModel):
    frequency: int

@router.post("/agent/{id}/frequency")
async def set_frequency(id: int, payload: dict):
    new_freq = payload.get("frequency")
    write_crawler_config(frequency=new_freq)
    await broadcast_status()
    return {"success": True}

@router.get("/monitor/status")
async def monitor_status():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT filename, url, status, last_checked FROM files ORDER BY last_checked DESC LIMIT 1")
    row = c.fetchone()
    conn.close()
    cfg = read_crawler_config()
    return [{
        "filename": row[0] if row else None,
        "url": row[1] if row else None,
        "status": row[2] if row else None,
        "last_checked": row[3] if row else None,
        "frequency": cfg["crawler_frequency"]
    }]

@router.get("/monitor/log")
def get_latest_log():
    return {"message": LATEST_STATUS}

@router.websocket("/ws/monitor")
async def websocket_endpoint(websocket: WebSocket):
    await manager.connect(websocket)
    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        manager.disconnect(websocket)

# ---------------------------
# Cleanup function
# ---------------------------
async def cleanup_mcp_client():
    """Clean up MCP client when router is being shut down"""
    global mcp_client
    if mcp_client:
        await mcp_client.__aexit__(None, None, None)
        mcp_client = None
